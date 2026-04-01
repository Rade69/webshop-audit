"""
Audit Report Generator — generiše Word (.docx) izvještaj iz rezultata audit run-a.

Može se koristiti standalone (CLI) ili iz GUI-a.
LLM pozivi koriste Gemini Flash (GEMINI_API_KEY) sa fallbackom na
DeepSeek (DEEPSEEK_API_KEY). Ako nijedan ključ nije dostupan,
LLM sekcije se zamjenjuju placeholderom.
"""
import json
import os
import warnings
from typing import Optional
from urllib.parse import urlparse

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


# ── Konstante ─────────────────────────────────────────────────────────────────

COLOR_DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
COLOR_DARK_GREY = RGBColor(0x40, 0x40, 0x40)
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GREEN_BG  = "E6F4EF"
COLOR_YELLOW_BG = "FDF3E3"
COLOR_RED_BG    = "FDEAEA"
COLOR_HEADER_BG = "1F3864"
COLOR_ALT_ROW   = "F2F7FF"
COLOR_BORDER    = "CCCCCC"


# ── LLM poziv ─────────────────────────────────────────────────────────────────

def call_llm(prompt: str, max_tokens: int = 500) -> str:
    """
    Poziva Gemini Flash za narativni tekst.
    Fallback na DeepSeek ako Gemini nije dostupan.
    Vraća placeholder string ako oba failaju.
    """
    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    deepseek_key = os.environ.get("DEEPSEEK_API_KEY", "")

    # Pokušaj Gemini
    if gemini_key:
        try:
            from google import genai
            client = genai.Client(api_key=gemini_key)
            response = client.models.generate_content(
                model="gemini-2.0-flash-exp",
                contents=prompt,
            )
            return response.text.strip()
        except Exception as e:
            warnings.warn(f"Gemini API poziv neuspješan: {e}. Pokušavam DeepSeek...")

    # Fallback DeepSeek
    if deepseek_key:
        try:
            import requests as _req
            resp = _req.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers={"Authorization": f"Bearer {deepseek_key}"},
                json={
                    "model": "deepseek-chat",
                    "messages": [{"role": "user", "content": prompt}],
                    "max_tokens": max_tokens,
                },
                timeout=30,
            )
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"].strip()
        except Exception as e:
            warnings.warn(f"DeepSeek API poziv neuspješan: {e}")

    return "[LLM generation failed — please add narrative manually]"


# ── Pomoćne funkcije za docx ───────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Postavlja background boju ćelije."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell):
    """Postavlja tanke bordere na ćeliju."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), COLOR_BORDER)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_heading(doc: Document, text: str, level: int):
    """Dodaje heading sa prilagođenim stilom."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_DARK_BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(9)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_DARK_BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_DARK_GREY
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(5)
    return p


def _add_table_header(table, headers: list[str]):
    """Puni header red tabele — tamno plava pozadina, bijeli tekst."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = COLOR_WHITE
        run.font.size = Pt(10)
        _set_cell_bg(cell, COLOR_HEADER_BG)
        _set_cell_border(cell)


def _add_table_row(table, values: list[str], alt: bool = False, score_col: int = -1):
    """Dodaje red u tabelu sa opcionim alternating bojama i score bojanjem."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        _set_cell_border(cell)

        if i == score_col and val not in ("-", ""):
            try:
                score = float(str(val).replace("/100", ""))
                if score >= 70:
                    _set_cell_bg(cell, COLOR_GREEN_BG)
                elif score >= 40:
                    _set_cell_bg(cell, COLOR_YELLOW_BG)
                else:
                    _set_cell_bg(cell, COLOR_RED_BG)
            except ValueError:
                pass
        elif alt:
            _set_cell_bg(cell, COLOR_ALT_ROW)


def _score_color(score: float) -> str:
    if score >= 70:
        return COLOR_GREEN_BG
    elif score >= 40:
        return COLOR_YELLOW_BG
    return COLOR_RED_BG


def _pct(count: int, total: int) -> str:
    if total == 0:
        return "0%"
    return f"{round(count / total * 100)}%"


def _domain_from_data(summary: dict, df_scored) -> str:
    """
    Izvlači domain iz podataka.
    Prioritet: summary fields → prvi URL iz products_scored.csv
    """
    for key in ("domain", "sitemap_url", "base_url"):
        val = summary.get(key, "")
        if val:
            parsed = urlparse(val)
            return parsed.netloc or val

    # Fallback: uzmi domain iz prvog URL-a u scored CSV-u
    if df_scored is not None and not df_scored.empty and "url" in df_scored.columns:
        first_url = str(df_scored["url"].dropna().iloc[0]) if len(df_scored) > 0 else ""
        if first_url.startswith("http"):
            parsed = urlparse(first_url)
            return parsed.netloc

    return "nepoznat webshop"


def _detect_language(domain: str) -> str:
    """Vraća 'bs' za .ba/.hr/.rs domene, inače 'en'."""
    for tld in (".ba", ".hr", ".rs", ".me", ".si"):
        if domain.endswith(tld):
            return "bs"
    return "en"


# ── Učitavanje podataka ────────────────────────────────────────────────────────

def _load_data(output_dir: str) -> dict:
    """Učitava sve relevantne fajlove iz output direktorija."""
    data = {}

    summary_path = os.path.join(output_dir, "run_summary.json")
    if os.path.isfile(summary_path):
        with open(summary_path, encoding="utf-8") as f:
            data["summary"] = json.load(f)
    else:
        data["summary"] = {}

    for name in ("products_scored", "manual_review_candidates", "category_summary", "errors"):
        path = os.path.join(output_dir, f"{name}.csv")
        if os.path.isfile(path):
            try:
                data[name] = pd.read_csv(path)
            except Exception:
                data[name] = None
        else:
            data[name] = None

    return data


def _compute_stats(data: dict) -> dict:
    """Računa sve agregatne statistike potrebne za izvještaj."""
    summary = data.get("summary", {})
    df = data.get("products_scored")

    total = summary.get("total_urls", 0)
    parsed = summary.get("successfully_parsed", 0)
    errors = summary.get("errors", 0)

    if df is not None and not df.empty:
        avg_overall  = round(df["overall_score"].mean(), 1) if "overall_score" in df.columns else 0
        avg_catalog  = round(df["catalog_score"].mean(), 1) if "catalog_score" in df.columns else 0
        avg_machine  = round(df["machine_score"].mean(), 1) if "machine_score" in df.columns else 0
        avg_commerce = round(df["commerce_score"].mean(), 1) if "commerce_score" in df.columns else 0

        no_schema    = int(df["schema_product_present"].eq(False).sum()) if "schema_product_present" in df.columns else 0
        no_offer     = int(df["schema_offer_present"].eq(False).sum()) if "schema_offer_present" in df.columns else 0
        no_price_sc  = int(df["suspicious_price_missing"].sum()) if "suspicious_price_missing" in df.columns else 0
        no_avail     = int(df["schema_availability"].isna().sum()) if "schema_availability" in df.columns else 0
        no_sku       = int(df["schema_sku"].isna().sum()) if "schema_sku" in df.columns else 0
        no_brand     = int(df["schema_brand"].isna().sum()) if "schema_brand" in df.columns else 0
        no_h1        = int(df["h1"].isna().sum()) if "h1" in df.columns else 0
        no_meta      = int(df["meta_description"].isna().sum()) if "meta_description" in df.columns else 0
        low_content  = int(df["suspicious_low_content"].sum()) if "suspicious_low_content" in df.columns else 0
        no_price_html = int(df["html_price_text"].isna().sum()) if "html_price_text" in df.columns else 0
        no_shipping  = int(df["shipping_signal"].eq(False).sum()) if "shipping_signal" in df.columns else 0
        no_returns   = int(df["returns_signal"].eq(False).sum()) if "returns_signal" in df.columns else 0
        # Handle NaN values in indexability_flags before using .str accessor
        if "indexability_flags" in df.columns:
            # Fill NaN with empty string for string operations
            flags_series = df["indexability_flags"].fillna("")
            noindex = int(flags_series.str.contains("noindex", na=False).sum())
            canonical_mm = int(flags_series.str.contains("canonical", na=False).sum())
        else:
            noindex = 0
            canonical_mm = 0
    else:
        avg_overall = avg_catalog = avg_machine = avg_commerce = 0
        no_schema = no_offer = no_price_sc = no_avail = no_sku = no_brand = 0
        no_h1 = no_meta = low_content = no_price_html = no_shipping = no_returns = 0
        noindex = canonical_mm = 0

    schema_pct      = _pct(summary.get("pages_with_schema_product", 0), parsed)
    price_miss_pct  = _pct(summary.get("pages_without_price", 0), parsed)
    low_cont_pct    = _pct(summary.get("pages_with_low_content", 0), parsed)

    # Biggest problems za 30/60/90 plan
    problems = sorted([
        ("Nema Product schema", no_schema, _pct(no_schema, parsed)),
        ("Nema cijene u structured data", no_price_sc, _pct(no_price_sc, parsed)),
        ("Malo vidljivog teksta", low_content, _pct(low_content, parsed)),
        ("Nema H1", no_h1, _pct(no_h1, parsed)),
        ("Noindex stranice", noindex, _pct(noindex, parsed)),
    ], key=lambda x: x[1], reverse=True)

    return {
        "total": total, "parsed": parsed, "errors": errors,
        "avg_overall": avg_overall, "avg_catalog": avg_catalog,
        "avg_machine": avg_machine, "avg_commerce": avg_commerce,
        "schema_pct": schema_pct, "price_miss_pct": price_miss_pct,
        "low_cont_pct": low_cont_pct,
        "no_schema": no_schema, "no_offer": no_offer,
        "no_price_sc": no_price_sc, "no_avail": no_avail,
        "no_sku": no_sku, "no_brand": no_brand,
        "no_h1": no_h1, "no_meta": no_meta,
        "low_content": low_content, "no_price_html": no_price_html,
        "no_shipping": no_shipping, "no_returns": no_returns,
        "noindex": noindex, "canonical_mm": canonical_mm,
        "candidates": summary.get("manual_review_candidates", 0),
        "problems": problems,
    }


# ── LLM promptovi ─────────────────────────────────────────────────────────────

def _prompt_executive_summary(domain: str, s: dict, lang: str) -> str:
    lang_instr = (
        "Write in Bosnian/Croatian/Serbian Latin script."
        if lang == "bs"
        else "Write in English."
    )
    return f"""You are writing an executive summary for a webshop audit report.
Write 3-4 concise paragraphs in a professional tone. {lang_instr}

Data about the shop:
- Domain: {domain}
- Total product pages scanned: {s['total']}
- Successfully parsed: {s['parsed']}
- Average overall score: {s['avg_overall']}/100
- Average catalog score: {s['avg_catalog']}/100
- Average machine score: {s['avg_machine']}/100
- Average commerce score: {s['avg_commerce']}/100
- Pages with Product schema: {s['total'] - s['no_schema']} ({s['schema_pct']})
- Pages missing price: {s['no_price_sc']} ({s['price_miss_pct']})
- Pages with low content: {s['low_content']} ({s['low_cont_pct']})
- Manual review candidates: {s['candidates']}
- Errors during fetch: {s['errors']}

Cover these points:
1. Overall impression of the shop's data quality
2. Most critical finding (the biggest single problem)
3. What this means for AI agent discoverability specifically
4. General direction for improvement

Be direct and specific. Avoid generic phrases.
Do not use bullet points — write in paragraphs.
Maximum 250 words."""


def _prompt_quick_wins(s: dict, domain: str, lang: str) -> str:
    lang_instr = (
        "Write in Bosnian/Croatian/Serbian Latin script."
        if lang == "bs"
        else "Write in English."
    )
    p = s["problems"]
    return f"""Based on this webshop audit data, identify the top 3 quick wins — specific, actionable improvements that would have the highest impact with the least effort. {lang_instr}

Data:
- Pages missing Product schema: {s['no_schema']} ({_pct(s['no_schema'], s['parsed'])})
- Pages missing price in schema: {s['no_price_sc']} ({_pct(s['no_price_sc'], s['parsed'])})
- Pages missing H1: {s['no_h1']}
- Pages missing meta description: {s['no_meta']}
- Pages with low content: {s['low_content']}
- Pages with canonical mismatch: {s['canonical_mm']}
- Domain: {domain}

Write exactly 3 quick wins.
Format each as:
**[Short title]**
[2-3 sentences: what to fix, why it matters for AI agent discoverability, expected impact]"""


def _prompt_plan(s: dict, domain: str, lang: str) -> str:
    lang_instr = (
        "Write in Bosnian/Croatian/Serbian Latin script."
        if lang == "bs"
        else "Write in English."
    )
    p = s["problems"]
    biggest = p[0][0] if p else "Nema podataka"
    biggest_pct = p[0][2] if p else "0%"
    second = p[1][0] if len(p) > 1 else "-"
    third = p[2][0] if len(p) > 2 else "-"
    return f"""Create a practical 30/60/90 day improvement plan for this webshop based on the audit findings. {lang_instr}

Key findings:
- Biggest problem: {biggest} (affects {biggest_pct} of pages)
- Second problem: {second}
- Third problem: {third}
- Overall score: {s['avg_overall']}/100
- Machine readability score: {s['avg_machine']}/100
- Domain: {domain}

Structure the plan as three phases:

**Prvih 30 dana — Kritični popravci**
[3-4 bullet points: things to fix immediately, highest impact]

**31-60 dana — Sistemski popravci**
[3-4 bullet points: structural improvements]

**61-90 dana — Optimizacija i provjera**
[3-4 bullet points: monitoring, testing, iterating]

Be specific and practical."""


# ── Gradnja dokumenta ──────────────────────────────────────────────────────────

def generate_report(output_dir: str, log=None) -> str:
    """
    Generiše Word izvještaj iz output_dir i snima ga kao audit_report.docx.

    Args:
        output_dir: Putanja do direktorija sa rezultatima audit run-a.
        log: Opcionalna funkcija (level, message) za log poruke.

    Returns:
        Putanja do generisanog .docx fajla.
    """
    def _log(level: str, msg: str):
        if log:
            log(level, msg)
        else:
            print(f"[{level.upper()}] {msg}")

    _log("info", f"Učitavam podatke iz: {output_dir}")
    data = _load_data(output_dir)
    s = _compute_stats(data)
    summary = data["summary"]
    domain = _domain_from_data(summary, data.get("products_scored"))
    lang = _detect_language(domain)
    timestamp = summary.get("timestamp", "")

    _log("info", "Pozivam LLM za narativne sekcije...")
    exec_summary = call_llm(_prompt_executive_summary(domain, s, lang), max_tokens=500)
    _log("info", "Executive summary generisan.")
    quick_wins = call_llm(_prompt_quick_wins(s, domain, lang), max_tokens=400)
    _log("info", "Quick wins generisani.")
    plan_text = call_llm(_prompt_plan(s, domain, lang), max_tokens=500)
    _log("info", "30/60/90 plan generisan.")

    doc = Document()

    # Postavi A4 format
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default font Arial
    for style in doc.styles:
        try:
            style.font.name = "Arial"
        except Exception:
            pass

    # ── Sekcija 1: Naslovna strana ─────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("WEBSHOP AUDIT REPORT")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = COLOR_DARK_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Agent-Friendly Readiness Assessment")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_DARK_GREY

    doc.add_paragraph()

    for label, value in [
        ("Shop:", domain),
        ("Datum:", timestamp),
        ("Ukupno stranica:", str(s["total"])),
        ("Uspješno obrađeno:", str(s["parsed"])),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bold_run = p.add_run(f"{label} ")
        bold_run.bold = True
        bold_run.font.size = Pt(12)
        val_run = p.add_run(value)
        val_run.font.size = Pt(12)

    doc.add_page_break()

    # ── Sekcija 2: Executive Summary ───────────────────────────────────────────
    _add_heading(doc, "1. Executive Summary", 1)
    for para in exec_summary.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(8)

    # ── Sekcija 3: Scorecard ───────────────────────────────────────────────────
    _add_heading(doc, "2. Scorecard", 1)

    score_table = doc.add_table(rows=1, cols=3)
    score_table.style = "Table Grid"
    _add_table_header(score_table, ["Dimenzija", "Prosječan score", "Što mjeri"])

    score_rows = [
        ("Catalog Score", f"{s['avg_catalog']}/100", "HTML kompletnost — naslov, opis, slike, tekst"),
        ("Machine Score", f"{s['avg_machine']}/100", "Schema.org structured data (JSON-LD)"),
        ("Commerce Score", f"{s['avg_commerce']}/100", "Cijena, valuta, dostupnost, SKU, GTIN"),
        ("Overall Score", f"{s['avg_overall']}/100", "Ponderisani prosjek sva tri"),
    ]
    for i, (dim, score, desc) in enumerate(score_rows):
        row = score_table.add_row()
        row.cells[0].text = dim
        row.cells[1].text = score
        row.cells[2].text = desc
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            _set_cell_border(cell)
            if i % 2 == 1:
                _set_cell_bg(cell, COLOR_ALT_ROW)
        # Oboji score ćeliju
        try:
            score_val = float(score.replace("/100", ""))
            _set_cell_bg(row.cells[1], _score_color(score_val))
            row.cells[1].paragraphs[0].runs[0].bold = True
        except ValueError:
            pass

    doc.add_paragraph()

    # Ključne metrike
    _add_heading(doc, "Ključne metrike", 3)
    metrics = [
        ("Stranica sa Product schema:", f"{s['total'] - s['no_schema']} od {s['parsed']} ({s['schema_pct']})"),
        ("Stranica bez cijene:", f"{s['no_price_sc']} ({s['price_miss_pct']})"),
        ("Stranica sa malo sadržaja:", f"{s['low_content']} ({s['low_cont_pct']})"),
        ("Stranica sa indexability problemima:", str(s["noindex"] + s["canonical_mm"])),
        ("Greške pri fetch-u:", str(s["errors"])),
    ]
    for label, value in metrics:
        p = doc.add_paragraph(style="List Bullet")
        run_l = p.add_run(label + " ")
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_v = p.add_run(value)
        run_v.font.size = Pt(11)

    # ── Sekcija 4: Ključni nalazi ──────────────────────────────────────────────
    _add_heading(doc, "3. Ključni nalazi", 1)

    # 4.1 Schema
    _add_heading(doc, "3.1 Schema i structured data", 2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    _add_table_header(t, ["Nalaz", "Broj stranica", "% od ukupno"])
    schema_rows = [
        ("Nema Product schema", s["no_schema"], _pct(s["no_schema"], s["parsed"])),
        ("Nema Offer schema", s["no_offer"], _pct(s["no_offer"], s["parsed"])),
        ("Nema cijene u schema", s["no_price_sc"], _pct(s["no_price_sc"], s["parsed"])),
        ("Nema availability", s["no_avail"], _pct(s["no_avail"], s["parsed"])),
        ("Nema SKU", s["no_sku"], _pct(s["no_sku"], s["parsed"])),
        ("Nema brand", s["no_brand"], _pct(s["no_brand"], s["parsed"])),
    ]
    for i, (label, count, pct) in enumerate(schema_rows):
        _add_table_row(t, [label, str(count), pct], alt=(i % 2 == 1))

    doc.add_paragraph()

    # 4.2 HTML sadržaj
    _add_heading(doc, "3.2 HTML sadržaj i kompletnost", 2)
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = "Table Grid"
    _add_table_header(t2, ["Nalaz", "Broj stranica", "% od ukupno"])
    html_rows = [
        ("Nema H1", s["no_h1"], _pct(s["no_h1"], s["parsed"])),
        ("Nema meta description", s["no_meta"], _pct(s["no_meta"], s["parsed"])),
        ("Malo vidljivog teksta (<200 znakova)", s["low_content"], s["low_cont_pct"]),
        ("Nema HTML price signala", s["no_price_html"], _pct(s["no_price_html"], s["parsed"])),
        ("Nema shipping signala", s["no_shipping"], _pct(s["no_shipping"], s["parsed"])),
        ("Nema returns signala", s["no_returns"], _pct(s["no_returns"], s["parsed"])),
    ]
    for i, (label, count, pct) in enumerate(html_rows):
        _add_table_row(t2, [label, str(count), pct], alt=(i % 2 == 1))

    doc.add_paragraph()

    # 4.3 Indexability
    _add_heading(doc, "3.3 Indexability i tehnički problemi", 2)
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    _add_table_header(t3, ["Nalaz", "Broj stranica"])
    index_rows = [
        ("Noindex stranice", s["noindex"]),
        ("Canonical mismatch", s["canonical_mm"]),
        ("Fetch greške (4xx/5xx/timeout)", s["errors"]),
    ]
    for i, (label, count) in enumerate(index_rows):
        _add_table_row(t3, [label, str(count)], alt=(i % 2 == 1))

    doc.add_paragraph()

    # 4.4 Kategorije
    cat_df = data.get("category_summary")
    if cat_df is not None and not cat_df.empty:
        _add_heading(doc, "3.4 Kategorije — bottom 5 po ocjeni", 2)
        t4 = doc.add_table(rows=1, cols=5)
        t4.style = "Table Grid"
        _add_table_header(t4, ["Kategorija", "Proizvoda", "Avg score", "Bez schema", "Bez cijene"])
        sort_col = "avg_overall_score" if "avg_overall_score" in cat_df.columns else cat_df.columns[0]
        bottom5 = cat_df.nsmallest(5, sort_col) if sort_col in cat_df.columns else cat_df.head(5)
        for i, row in bottom5.iterrows():
            vals = [
                str(row.get("category", "-")),
                str(row.get("count", "-")),
                str(round(row.get("avg_overall_score", 0), 1)),
                str(row.get("no_schema_count", "-")),
                str(row.get("no_price_count", "-")),
            ]
            _add_table_row(t4, vals, alt=(i % 2 == 1), score_col=2)
        doc.add_paragraph()

    # ── Sekcija 5: Quick Wins ──────────────────────────────────────────────────
    _add_heading(doc, "4. Quick Wins", 1)
    for para in quick_wins.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith("**") and para.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(para.strip("*"))
            run.bold = True
            run.font.size = Pt(12)
        else:
            # Inline bold (**text**)
            p = doc.add_paragraph()
            parts = para.split("**")
            for j, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.bold = (j % 2 == 1)
                    run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)

    # ── Sekcija 6: 30/60/90 plan ───────────────────────────────────────────────
    _add_heading(doc, "5. Plan poboljšanja — 30/60/90 dana", 1)
    for para in plan_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith("**") and "**" in para[2:]:
            end = para.index("**", 2)
            heading_text = para[2:end]
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_DARK_BLUE
        elif para.startswith("- ") or para.startswith("* "):
            for line in para.split("\n"):
                line = line.lstrip("- *").strip()
                if line:
                    p = doc.add_paragraph(line, style="List Bullet")
                    p.runs[0].font.size = Pt(11)
        else:
            p = doc.add_paragraph(para)
            p.runs[0].font.size = Pt(11) if p.runs else None

    # ── Sekcija 7: Top 20 kandidata ───────────────────────────────────────────
    cand_df = data.get("manual_review_candidates")
    if cand_df is not None and not cand_df.empty:
        _add_heading(doc, "6. Prilog — Top 20 kandidata za pregled", 1)
        t5 = doc.add_table(rows=1, cols=4)
        t5.style = "Table Grid"
        _add_table_header(t5, ["URL", "Ukupna ocjena", "Razlog", "Oznake"])
        def _clean(val) -> str:
            """Pretvara pandas NaN i None u '-'."""
            if val is None:
                return "-"
            s = str(val).strip()
            return "-" if s.lower() in ("nan", "none", "") else s

        for i, row in cand_df.head(20).iterrows():
            url = _clean(row.get("url", ""))
            if len(url) > 60:
                url = url[:57] + "..."
            score_raw = row.get("overall_score", None) if "overall_score" in cand_df.columns else None
            try:
                score = str(round(float(score_raw), 1)) if score_raw is not None and str(score_raw).lower() != "nan" else "-"
            except (ValueError, TypeError):
                score = "-"
            missing = _clean(row.get("missing_fields", "")) if "missing_fields" in cand_df.columns else "-"
            flags   = _clean(row.get("indexability_flags", "")) if "indexability_flags" in cand_df.columns else "-"
            # Human-readable reason
            parts = []
            if missing != "-":
                if "schema" in missing.lower():
                    parts.append("Missing Schema")
                if "price" in missing.lower():
                    parts.append("Missing Price")
                if not parts:
                    parts.append(missing[:40])
            reason = ", ".join(parts) if parts else "-"
            _add_table_row(t5, [url, score, reason, flags[:40]], alt=(i % 2 == 1), score_col=1)

    # ── Snimanje ───────────────────────────────────────────────────────────────
    output_path = os.path.join(output_dir, "audit_report.docx")
    doc.save(output_path)
    _log("info", f"Izvještaj sačuvan: {output_path}")
    return output_path
